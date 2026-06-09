#!/usr/bin/env python3
"""
One-time script to generate the initial SafeVixAI_MASTER.docx
with Part A (static docs), Part B (semi-static config), and Part C marker.
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DOCS_DIR = os.path.join(ROOT, "docs")
OUTPUT = os.path.join(DOCS_DIR, "SafeVixAI_MASTER.docx")


def add_heading(doc, text, level, color_hex="0D1117"):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    return para


def add_para(doc, text, bold=False, italic=False, font_size=10, color_hex=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    return para


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def read_md(filepath):
    """Read a markdown file and return its content."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"[Error reading {filepath}: {e}]"


def md_to_docx_section(doc, title, md_content, heading_level=2):
    """
    Convert markdown content to docx paragraphs.
    Handles: headings, bullet points, code blocks, tables, and plain text.
    """
    add_heading(doc, title, heading_level, "1A5C38")

    in_code_block = False
    code_lines = []

    for line in md_content.split("\n"):
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block — flush
                code_text = "\n".join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.size = Pt(8)
                run.font.name = "Consolas"
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                para.paragraph_format.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip empty lines
        if not stripped:
            continue

        # Headings (### → level 4, ## → level 3)
        if stripped.startswith("####"):
            add_heading(doc, stripped.lstrip("#").strip(), min(heading_level + 3, 5))
        elif stripped.startswith("###"):
            add_heading(doc, stripped.lstrip("#").strip(), min(heading_level + 2, 4))
        elif stripped.startswith("##"):
            add_heading(doc, stripped.lstrip("#").strip(), min(heading_level + 1, 3))
        elif stripped.startswith("#"):
            # Skip top-level heading (already used as section title)
            continue
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        # Table rows (basic — just add as text)
        elif stripped.startswith("|"):
            # Skip separator rows
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                doc.add_paragraph("  |  ".join(cells), style="List Bullet")
        # Plain text
        else:
            add_para(doc, stripped, font_size=10)


def build_master_doc():
    doc = Document()

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_before = Pt(120)
    run = title_para.add_run("SafeVixAI")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x38)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Powered Road Safety Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x58, 0x58, 0x58)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        "Enterprise Master Document\n"
        "IIT Madras Road Safety Hackathon 2026"
    )
    run.font.size = Pt(14)
    run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.space_before = Pt(40)
    run = meta.add_run(
        "safevixai.vercel.app  •  github.com/SafeVixAI/SafeVixAI\n"
        "Structure: Part A (Static) + Part B (Semi-Static) + Part C (Live Auto-Updated)"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", 1, "0D1117")
    toc_items = [
        "PART A — STATIC DOCUMENTATION",
        "  A.1  Product Requirements (PRD)",
        "  A.2  System Architecture",
        "  A.3  Features & Modules",
        "  A.4  Technology Stack",
        "  A.5  API Reference (28 endpoints)",
        "  A.6  Database Schema",
        "  A.7  AI & Chatbot Architecture",
        "  A.8  Agent Tools & Intent Detection",
        "  A.9  Security & Authentication",
        "  A.10 Offline Architecture",
        "  A.11 UI/UX Design System (DESIGN.md)",
        "  A.12 Data Sources & Datasets",
        "  A.13 Deployment Guide",
        "  A.14 Setup & Development (SETUP.md)",
        "  A.15 Contributing Guidelines",
        "  A.16 Roadmap",
        "  A.17 Agent System Architecture (AGENTS.md)",
        "  A.18 Project README",
        "  A.19 Platform Capabilities (SKILL.md)",
        "  A.20 UI/UX Component Reference",
        "  A.21 Complete Resource Checklist",
        "",
        "PART B — SEMI-STATIC CONFIGURATION",
        "  B.1  Environment Variables (3 services)",
        "  B.2  Database Tables & Schema",
        "  B.3  Dataset Placement Guide",
        "",
        "PART C — LIVE STATUS (AUTO-UPDATED DAILY)",
        "  C.1  Repository Overview",
        "  C.2  Deployment & Service Health",
        "  C.3  Recent CI/CD Runs",
        "  C.4  Open GitHub Issues",
        "  C.5  Recent Commits",
        "  C.6  Feature Completion Status",
    ]
    for item in toc_items:
        if not item:
            doc.add_paragraph("")
            continue
        if item.startswith("PART"):
            add_para(doc, item, bold=True, font_size=12, color_hex="1A5C38")
        else:
            add_para(doc, item, font_size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # PART A — STATIC DOCUMENTATION
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART A — STATIC DOCUMENTATION", 1, "1A5C38")
    add_para(
        doc,
        "This section consolidates all project documentation into a single reference. "
        "Content is written once and updated only when architecture or design decisions change.",
        italic=True,
        font_size=10,
    )

    # Ordered list of docs to merge
    docs_to_merge = [
        ("A.1  Product Requirements (PRD)", "PRD.md"),
        ("A.2  System Architecture", "Architecture.md"),
        ("A.3  Features & Modules", "Features.md"),
        ("A.4  Technology Stack", "TechStack.md"),
        ("A.5  API Reference", "API.md"),
        ("A.6  Database Schema", "Database.md"),
        ("A.7  AI & Chatbot Architecture", "AI_Instructions.md"),
        ("A.8  Agent Tools & Intent Detection", "Agent.md"),
        ("A.9  Security & Authentication", "Security.md"),
        ("A.10 Offline Architecture", "Offline_Architecture.md"),
        ("A.11 UI/UX Design System", None),  # DESIGN.md is top-level
        ("A.12 Data Sources & Datasets", "DataSources.md"),
        ("A.13 Deployment Guide", "Deployment.md"),
        ("A.14 Setup & Development", None),  # SETUP.md is top-level
        ("A.15 Contributing Guidelines", "Contributing.md"),
        ("A.16 Roadmap", "Roadmap.md"),
        ("A.17 Agent System Architecture", None),  # AGENTS.md is top-level
        ("A.18 Project README", None),  # README.md is top-level
        ("A.19 Platform Capabilities", None),  # SKILL.md is top-level
        ("A.20 UI/UX Component Reference", "UIUX.md"),
        ("A.21 Complete Resource Checklist", "Complete_Project_Resource_Checklist.md"),
    ]

    for title, filename in docs_to_merge:
        doc.add_page_break()

        if filename:
            filepath = os.path.join(DOCS_DIR, filename)
        elif "Design" in title:
            filepath = os.path.join(ROOT, "DESIGN.md")
        elif "Setup" in title:
            filepath = os.path.join(ROOT, "SETUP.md")
        elif "Agent System" in title:
            filepath = os.path.join(ROOT, "AGENTS.md")
        elif "README" in title:
            filepath = os.path.join(ROOT, "README.md")
        elif "Capabilities" in title:
            filepath = os.path.join(ROOT, "SKILL.md")
        else:
            filepath = None

        if filepath and os.path.exists(filepath):
            content = read_md(filepath)
            md_to_docx_section(doc, title, content, heading_level=2)
            add_para(
                doc,
                f"— Source: {os.path.relpath(filepath, ROOT)} —",
                italic=True,
                font_size=8,
                color_hex="999999",
            )
        else:
            add_heading(doc, title, 2, "1A5C38")
            add_para(
                doc,
                f"[Document not found: {filename or 'N/A'}]",
                italic=True,
                color_hex="CC0000",
            )

    # ═══════════════════════════════════════════════════════════════════════
    # PART B — SEMI-STATIC CONFIGURATION
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "PART B — SEMI-STATIC CONFIGURATION", 1, "1A5C38")
    add_para(
        doc,
        "Configuration that changes only when the system architecture is modified. "
        "Update this section when adding new services, tables, or environment variables.",
        italic=True,
        font_size=10,
    )

    # B.1 — Environment Variables
    doc.add_page_break()
    env_file = os.path.join(DOCS_DIR, "Environment.md")
    if os.path.exists(env_file):
        md_to_docx_section(doc, "B.1  Environment Variables", read_md(env_file))
    else:
        add_heading(doc, "B.1  Environment Variables", 2, "1A5C38")

        # Backend
        add_heading(doc, "Backend Service (.env)", 3)
        backend_env = os.path.join(ROOT, "backend", ".env.example")
        if os.path.exists(backend_env):
            content = read_md(backend_env)
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.size = Pt(8)
            run.font.name = "Consolas"

        # Chatbot
        add_heading(doc, "Chatbot Service (.env)", 3)
        chatbot_env = os.path.join(ROOT, "chatbot_service", ".env.example")
        if os.path.exists(chatbot_env):
            content = read_md(chatbot_env)
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.size = Pt(8)
            run.font.name = "Consolas"

        # Frontend
        add_heading(doc, "Frontend (.env.local)", 3)
        frontend_env = os.path.join(ROOT, "frontend", ".env.example")
        if os.path.exists(frontend_env):
            content = read_md(frontend_env)
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.size = Pt(8)
            run.font.name = "Consolas"

    # B.2 — Database Schema
    doc.add_page_break()
    db_file = os.path.join(DOCS_DIR, "Database.md")
    if os.path.exists(db_file):
        md_to_docx_section(doc, "B.2  Database Tables & Schema", read_md(db_file))
    else:
        add_heading(doc, "B.2  Database Tables & Schema", 2, "1A5C38")
        add_para(doc, "[Database.md not found]", italic=True)

    # B.3 — Dataset Placement
    doc.add_page_break()
    dataset_file = os.path.join(DOCS_DIR, "DATASET_PLACEMENT.md")
    if os.path.exists(dataset_file):
        md_to_docx_section(
            doc, "B.3  Dataset Placement Guide", read_md(dataset_file)
        )
    else:
        add_heading(doc, "B.3  Dataset Placement Guide", 2, "1A5C38")
        add_para(doc, "[DATASET_PLACEMENT.md not found]", italic=True)

    # ═══════════════════════════════════════════════════════════════════════
    # PART C — LIVE STATUS MARKER
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "PART C — LIVE STATUS (AUTO-UPDATED DAILY)", 1, "1A5C38")
    add_para(
        doc,
        "⏳ This section will be populated automatically by the GitHub Actions workflow "
        "(scripts/update_master_doc.py). Run the workflow manually or wait for the daily "
        "9:00 AM IST scheduled run.\n\n"
        "To trigger manually: GitHub → Actions → Update Master Document → Run workflow",
        italic=True,
        font_size=10,
    )

    # ── Save ────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    file_size_kb = os.path.getsize(OUTPUT) // 1024
    print(f"✅ Master document created: {OUTPUT}")
    print(f"   Size: {file_size_kb} KB")
    print("   Parts: A (21 sections) + B (3 sections) + C (marker)")
    print("   Source: 18 docs/ files + 5 root files = 23 total")
    print("   Next: Run 'python scripts/update_master_doc.py' to populate Part C")


if __name__ == "__main__":
    build_master_doc()
