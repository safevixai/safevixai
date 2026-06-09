#!/usr/bin/env python3
"""
SafeVixAI — Living Master Document Auto-Updater
================================================
Triggered on push to docs/root .md files via GitHub Actions.
Fetches live data from GitHub API, Render, and Vercel,
then rewrites PART C of docs/SafeVixAI_MASTER.docx.

Usage:
    python scripts/update_master_doc.py

Environment:
    GITHUB_TOKEN  — GitHub Personal Access Token (auto-provided by Actions)
"""

import logging
import os
import sys
import requests
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# Inject project root to sys.path to access alert_service singleton
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

logger = logging.getLogger(__name__)

try:
    from alert_service import get_alert_service
except Exception:
    logger.debug("alert_service not available, alerting disabled", exc_info=True)
    get_alert_service = None

# Fix Windows cp1252 encoding crashes with emoji print statements
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    logger.debug("Could not reconfigure stdout encoding", exc_info=True)

# ─── Configuration ──────────────────────────────────────────────────────────

GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN", "")
REPO = os.environ.get("GITHUB_REPOSITORY", "SafeVixAI/SafeVixAI")
GH_HEADERS = {
    "Accept": "application/vnd.github.v3+json",
}
if GITHUB_TOKEN:
    GH_HEADERS["Authorization"] = f"Bearer {GITHUB_TOKEN}"
IST = timezone(timedelta(hours=5, minutes=30))

SERVICES = {
    "Backend API": os.environ.get("BACKEND_HEALTH_URL", "https://safevixai-api.onrender.com/health"),
    "Chatbot Service": os.environ.get(
        "CHATBOT_HEALTH_URL",
        "https://safevixai-chatbot.onrender.com/health",
    ),
    "Frontend (Vercel)": os.environ.get("FRONTEND_URL", "https://safevixai.vercel.app"),
}

MASTER_DOC_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "docs",
    "SafeVixAI_MASTER.docx",
)

# ─── Data Fetchers ──────────────────────────────────────────────────────────


def fetch_open_issues():
    """Fetch all open GitHub issues (up to 50)."""
    try:
        r = requests.get(
            f"https://api.github.com/repos/{REPO}/issues",
            params={"state": "open", "per_page": 50},
            headers=GH_HEADERS,
            timeout=15,
        )
        if r.ok:
            # Filter out pull requests (they also show up as issues)
            return [i for i in r.json() if "pull_request" not in i]
        print(f"  ⚠ GitHub Issues API returned {r.status_code}")
        return []
    except Exception as e:
        print(f"  ⚠ GitHub Issues fetch failed: {e}")
        return []


def fetch_recent_commits(n=10):
    """Fetch the last N commits from the default branch."""
    try:
        r = requests.get(
            f"https://api.github.com/repos/{REPO}/commits",
            params={"per_page": n},
            headers=GH_HEADERS,
            timeout=15,
        )
        if not r.ok:
            print(f"  ⚠ GitHub Commits API returned {r.status_code}")
            return []
        return [
            {
                "sha": c["sha"][:7],
                "msg": c["commit"]["message"].split("\n")[0][:80],
                "author": c["commit"]["author"]["name"],
                "date": c["commit"]["author"]["date"][:10],
            }
            for c in r.json()
        ]
    except Exception as e:
        print(f"  ⚠ GitHub Commits fetch failed: {e}")
        return []


def fetch_service_health():
    """Ping each deployed service and record latency + status."""
    results = {}
    for name, url in SERVICES.items():
        try:
            r = requests.get(url, timeout=15)
            ms = int(r.elapsed.total_seconds() * 1000)
            # Try to extract version from JSON health response
            version = "?"
            content_type = r.headers.get("content-type", "")
            if "json" in content_type:
                try:
                    version = r.json().get("version", "?")
                except Exception:
                    logger.debug("Failed to parse version from health response", exc_info=True)
            results[name] = {
                "status": "UP",
                "ms": ms,
                "version": version,
                "code": r.status_code,
            }
        except requests.exceptions.Timeout:
            results[name] = {"status": "DOWN", "error": "Connection timed out (15s)"}
        except requests.exceptions.ConnectionError:
            results[name] = {"status": "DOWN", "error": "Connection refused / DNS failure"}
        except Exception as e:
            results[name] = {"status": "DOWN", "error": str(e)[:60]}

        if results[name]["status"] == "DOWN" and get_alert_service:
            try:
                get_alert_service().alert_external_api_failed(
                    service_name=f"Deployment ping ({name})",
                    endpoint=url,
                    status_code=0,
                    error_msg=results[name]["error"],
                )
            except Exception:
                logger.debug("Failed to dispatch alert for service %s", name, exc_info=True)
    return results


def fetch_workflow_runs():
    """Fetch the last 5 GitHub Actions workflow runs."""
    try:
        r = requests.get(
            f"https://api.github.com/repos/{REPO}/actions/runs",
            params={"per_page": 8},
            headers=GH_HEADERS,
            timeout=15,
        )
        if not r.ok:
            print(f"  ⚠ GitHub Actions API returned {r.status_code}")
            return []
        return [
            {
                "name": run["name"],
                "status": run["status"],
                "conclusion": run.get("conclusion", "in_progress"),
                "branch": run["head_branch"],
                "date": run["created_at"][:10],
            }
            for run in r.json().get("workflow_runs", [])[:8]
        ]
    except Exception as e:
        print(f"  ⚠ GitHub Actions fetch failed: {e}")
        return []


def fetch_repo_stats():
    """Fetch repository-level statistics (stars, forks, size)."""
    try:
        r = requests.get(
            f"https://api.github.com/repos/{REPO}",
            headers=GH_HEADERS,
            timeout=15,
        )
        if r.ok:
            data = r.json()
            return {
                "stars": data.get("stargazers_count", 0),
                "forks": data.get("forks_count", 0),
                "open_issues": data.get("open_issues_count", 0),
                "size_kb": data.get("size", 0),
                "default_branch": data.get("default_branch", "main"),
                "updated_at": data.get("updated_at", "")[:10],
            }
        return {}
    except Exception:
        return {}


# ─── DOCX Helpers ───────────────────────────────────────────────────────────


def add_colored_heading(doc, text, level, hex_color="1A5C38"):
    """Add a heading with a custom color."""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
    return para


def add_styled_para(doc, text, bold=False, italic=False, font_size=10):
    """Add a paragraph with optional styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return para


def add_status_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    return table


# ─── Main Update Logic ─────────────────────────────────────────────────────


def update_part_c(doc_path: str):
    """
    Open the master DOCX, find the PART C marker,
    delete everything after it, and write fresh live data.
    """
    if not os.path.exists(doc_path):
        print(f"✘ Master doc not found at: {doc_path}")
        sys.exit(1)

    doc = Document(doc_path)
    now = datetime.now(IST).strftime("%Y-%m-%d %H:%M IST")

    body = doc.element.body
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    sectPr = body.find(f"{ns}sectPr")

    # ── Find PART C marker and clear everything after it ────────────────
    part_c_idx = None
    # Iterate backwards so we hit the actual heading instead of the Table of Contents
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        if "PART C" in para.text and "LIVE" in para.text.upper() and para.style.name.startswith("Heading"):
            part_c_idx = i
            break

    if part_c_idx is None:
        print("⚠ PART C marker not found — appending at end")
    else:
        # Remove all elements after PART C heading, but preserve sectPr
        part_c_element = doc.paragraphs[part_c_idx]._element
        # Find all siblings after the PART C paragraph
        found = False
        elements_to_remove = []
        for child in body:
            if child is part_c_element:
                found = True
                continue
            if found:
                # Never remove sectPr — it holds page layout metadata
                if child.tag == f"{ns}sectPr":
                    continue
                elements_to_remove.append(child)
        for elem in elements_to_remove:
            body.remove(elem)

    print(f"📡 Fetching live data at {now}...")

    # ── WRITE FRESH PART C CONTENT ──────────────────────────────────────

    # Timestamp
    add_styled_para(doc, f"🕐 Last auto-updated: {now}", bold=True, font_size=11)
    add_styled_para(
        doc,
        "This section is automatically rewritten on push to docs/ or root .md files via GitHub Actions. "
        "Do not edit manually — changes will be overwritten.",
        italic=True,
        font_size=9,
    )

    # ── Section 1: Repository Overview ──────────────────────────────────
    add_colored_heading(doc, "Repository Overview", 2)
    stats = fetch_repo_stats()
    if stats:
        add_status_table(
            doc,
            ["Metric", "Value"],
            [
                ["⭐ Stars", stats.get("stars", "N/A")],
                ["🍴 Forks", stats.get("forks", "N/A")],
                ["📦 Repo Size", f"{stats.get('size_kb', 0) // 1024} MB"],
                ["🔀 Default Branch", stats.get("default_branch", "main")],
                ["📅 Last Updated", stats.get("updated_at", "N/A")],
                ["🐛 Open Issues", stats.get("open_issues", "N/A")],
            ],
        )
    else:
        add_styled_para(doc, "⚠ Could not fetch repository statistics.", italic=True)

    # ── Section 2: Service Health ───────────────────────────────────────
    add_colored_heading(doc, "Deployment & Service Health", 2)
    health = fetch_service_health()
    health_rows = []
    for name, h in health.items():
        if h["status"] == "UP":
            status = "✅ UP"
            detail = f"{h.get('ms', '?')}ms | HTTP {h.get('code', '?')} | v{h.get('version', '?')}"
        else:
            status = "🔴 DOWN"
            detail = h.get("error", "Unknown error")
        health_rows.append([name, status, detail])
    add_status_table(doc, ["Service", "Status", "Details"], health_rows)

    up_count = sum(1 for h in health.values() if h["status"] == "UP")
    total = len(health)
    add_styled_para(
        doc,
        f"Overall: {up_count}/{total} services operational.",
        bold=True,
        font_size=10,
    )

    # ── Section 3: CI/CD Pipeline Status ────────────────────────────────
    add_colored_heading(doc, "Recent CI/CD Runs", 2)
    runs = fetch_workflow_runs()
    if runs:
        ci_rows = []
        for run in runs:
            conclusion = run.get("conclusion") or run["status"]
            if conclusion == "success":
                icon = "✅"
            elif conclusion == "failure":
                icon = "🔴"
            elif conclusion == "cancelled":
                icon = "⚪"
            else:
                icon = "⏳"
            ci_rows.append(
                [
                    f"{icon} {run['name']}",
                    run["branch"],
                    conclusion,
                    run["date"],
                ]
            )
        add_status_table(doc, ["Workflow", "Branch", "Result", "Date"], ci_rows)
    else:
        add_styled_para(doc, "No recent CI/CD runs found.", italic=True)

    # ── Section 4: Open GitHub Issues ───────────────────────────────────
    add_colored_heading(doc, "Open GitHub Issues", 2)
    issues = fetch_open_issues()
    critical = [
        i
        for i in issues
        if any(
            l["name"].lower() in ["critical", "bug", "p0", "security"]
            for l in i.get("labels", [])
        )
    ]
    add_styled_para(
        doc,
        f"Total open: {len(issues)} | Critical/Bug: {len(critical)}",
        bold=True,
    )

    if issues:
        issue_rows = []
        for issue in issues[:25]:
            labels = ", ".join(l["name"] for l in issue.get("labels", []))
            assignee = (issue.get("assignee") or {}).get("login", "—")
            issue_rows.append(
                [
                    f"#{issue['number']}",
                    issue["title"][:60],
                    labels or "—",
                    assignee,
                ]
            )
        add_status_table(doc, ["#", "Title", "Labels", "Assignee"], issue_rows)
        if len(issues) > 25:
            add_styled_para(
                doc,
                f"... and {len(issues) - 25} more → github.com/{REPO}/issues",
                italic=True,
                font_size=9,
            )
    else:
        add_styled_para(doc, "🎉 No open issues — all clear!", font_size=10)

    # ── Section 5: Recent Commits ───────────────────────────────────────
    add_colored_heading(doc, "Recent Commits (last 10)", 2)
    commits = fetch_recent_commits(10)
    if commits:
        commit_rows = []
        for c in commits:
            commit_rows.append([c["sha"], c["date"], c["msg"], c["author"]])
        add_status_table(doc, ["SHA", "Date", "Message", "Author"], commit_rows)
    else:
        add_styled_para(doc, "No commits found.", italic=True)

    # ── Section 6: Feature Completion Matrix ────────────────────────────
    add_colored_heading(doc, "Feature Completion Status", 2)
    add_status_table(
        doc,
        ["Module", "Status", "Confidence"],
        [
            ["Emergency Locator (GPS + SOS)", "✅ Production", "95%"],
            ["AI Chatbot (11 LLM providers)", "✅ Production", "90%"],
            ["Challan Calculator (DuckDB)", "✅ Production", "95%"],
            ["Road Reporter (RoadWatch)", "✅ Production", "90%"],
            ["Crash Detection (DeviceMotion)", "✅ Production", "85%"],
            ["Offline AI (WebLLM Phi-3)", "✅ Production", "85%"],
            ["Live Family Tracking", "✅ Production", "80%"],
            ["Bystander Mode", "✅ Production", "80%"],
            ["PWA (Offline + Install)", "✅ Production", "90%"],
            ["Waze CIFS Feed", "✅ Production", "85%"],
        ],
    )

    # ── Section 7: Production Monitoring & Alerting ─────────────────────
    add_colored_heading(doc, "Production Monitoring & Alerting", 2)
    add_styled_para(
        doc,
        "SafeVixAI uses alert_service.py (project root) for production failure notifications. "
        "Email alerts are sent via Gmail SMTP when critical systems fail.",
        font_size=10,
    )
    add_status_table(
        doc,
        ["Service", "Monitored By", "Trigger"],
        [
            ["9 LLM Providers", "chatbot/providers/router.py", "All fallback providers fail"],
            ["Backend APIs (Overpass, Nominatim, OSRM)", "chatbot/tools/__init__.py", "HTTP 5xx or timeout"],
            ["PostgreSQL/PostGIS Database", "backend/main.py", "/health returns DB unavailable"],
            ["Unhandled Backend Errors", "backend/main.py", "Any unhandled 500 exception"],
            ["Wiki Doc Generation", "scripts/wiki_manager.py", "5+ consecutive LLM failures"],
        ],
    )
    add_styled_para(
        doc,
        "Each alert includes 3 diagnostic solutions + 5-min cooldown per alert type.",
        italic=True,
        font_size=9,
    )

    # ── Section 8: Wiki Documentation Stats ─────────────────────────────
    add_colored_heading(doc, "Auto-Generated Wiki Documentation", 2)
    wiki_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "docs", "wiki", "content")
    wiki_count = 0
    mermaid_count = 0
    if os.path.isdir(wiki_dir):
        for root_d, _, files in os.walk(wiki_dir):
            for f in files:
                if f.endswith(".md"):
                    wiki_count += 1
                    try:
                        text = open(os.path.join(root_d, f), encoding="utf-8").read()
                        mermaid_count += text.count("```mermaid")
                    except Exception:
                        logger.debug("Failed to read wiki file %s for mermaid count", f, exc_info=True)
    add_status_table(
        doc,
        ["Metric", "Value"],
        [
            ["Wiki Pages", str(wiki_count)],
            ["Mermaid Diagrams", str(mermaid_count)],
            ["Generation Method", "LLM (OpenRouter → Mistral → Gemini)"],
            ["CI Trigger", "Push to backend/chatbot/frontend/docs"],
            ["Source of Truth", "Code → LLM → Wiki (never manually edit)"],
        ],
    )

    # ── Footer ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_styled_para(
        doc,
        f"— End of Auto-Generated Section —\n"
        f"Generated by scripts/update_master_doc.py\n"
        f"Timestamp: {now}\n"
        f"Repository: github.com/{REPO}",
        italic=True,
        font_size=8,
    )

    # Ensure sectPr is the last element of the body (OpenXML spec compliance)
    sectPr_el = body.find(f"{ns}sectPr")
    if sectPr_el is not None:
        body.remove(sectPr_el)
        body.append(sectPr_el)
    elif sectPr is not None:
        body.append(sectPr)

    # ── Save ────────────────────────────────────────────────────────────
    doc.save(doc_path)
    print(f"✅ Master doc updated at {now}")
    print(f"   Services: {up_count}/{total} UP")
    print(f"   Open issues: {len(issues)}")
    print(f"   Recent commits: {len(commits)}")
    print(f"   CI runs: {len(runs)}")


# ─── Entry Point ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if not GITHUB_TOKEN:
        print("⚠ GITHUB_TOKEN not set — GitHub API calls will be rate-limited")
    update_part_c(MASTER_DOC_PATH)
